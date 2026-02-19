#!/usr/bin/env python3
"""
Generador de Documento de Proyecto Xcapit FHE-ML
Combina toda la documentación, diagramas SVG y demos capturados en un DOCX completo.
"""

import os
import io
import re
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

try:
    import cairosvg
    HAS_CAIROSVG = True
except ImportError:
    HAS_CAIROSVG = False
    print("Warning: cairosvg not available, SVGs will be converted with limited support")

# Paths
PROJECT_ROOT = Path(__file__).parent.parent
DOCS_DIR = PROJECT_ROOT / "docs"
DIAGRAMS_DIR = DOCS_DIR / "diagrams"
OUTPUT_DIR = PROJECT_ROOT / "output"
SCREENSHOTS_DIR = OUTPUT_DIR / "screenshots"


def setup_directories():
    """Create necessary directories."""
    OUTPUT_DIR.mkdir(exist_ok=True)
    SCREENSHOTS_DIR.mkdir(exist_ok=True)


def create_document_styles(doc):
    """Configure document styles."""
    # Title style
    style = doc.styles['Title']
    style.font.size = Pt(28)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # Heading 1
    style = doc.styles['Heading 1']
    style.font.size = Pt(18)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

    # Heading 2
    style = doc.styles['Heading 2']
    style.font.size = Pt(14)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)

    # Heading 3
    style = doc.styles['Heading 3']
    style.font.size = Pt(12)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    return doc


def add_cover_page(doc):
    """Add cover page."""
    for _ in range(3):
        doc.add_paragraph()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Xcapit FHE-ML Platform")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Plataforma de Machine Learning\ncon Encriptación Homomórfica Completa")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()
    doc.add_paragraph()

    # Description
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run(
        "Documentación Técnica Completa\n"
        "Fundamentos Teóricos, Arquitectura, Guías Prácticas y Demos"
    )
    run.font.size = Pt(14)
    run.font.italic = True

    for _ in range(5):
        doc.add_paragraph()

    # Date
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"Versión 1.0.0\nDiciembre 2024")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.add_page_break()


def add_table_of_contents(doc):
    """Add table of contents."""
    doc.add_heading("Tabla de Contenidos", level=1)

    toc_items = [
        ("1. Introducción", 1),
        ("2. Fundamentos Teóricos", 1),
        ("   2.1 Encriptación Homomórfica", 2),
        ("   2.2 El Esquema CKKS", 2),
        ("   2.3 Machine Learning sobre Datos Encriptados", 2),
        ("   2.4 Aproximaciones Polinomiales", 2),
        ("3. Arquitectura del Sistema", 1),
        ("4. Modelos de Machine Learning", 1),
        ("5. Guía de Inicio Rápido", 1),
        ("6. Demos Visuales de la Plataforma", 1),
        ("7. Diagramas del Sistema", 1),
        ("8. Glosario", 1),
        ("9. Referencias", 1),
    ]

    for item, level in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        if level == 1:
            run.font.bold = True
        run.font.size = Pt(11)

    doc.add_page_break()


def convert_svg_to_png(svg_path, output_path, width=700):
    """Convert SVG to PNG."""
    if not HAS_CAIROSVG:
        return None
    try:
        cairosvg.svg2png(url=str(svg_path), write_to=str(output_path), output_width=width)
        return output_path
    except Exception as e:
        print(f"  Error converting {svg_path.name}: {e}")
        return None


def add_styled_table(doc, headers, rows):
    """Create a styled Word table from data."""
    if not headers or not rows:
        return

    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row with styling
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header.strip().replace('**', '')
        # Style header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Header background color
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="4F46E5"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            clean_text = cell_text.strip().replace('**', '')
            cell.text = clean_text
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            # Alternate row colors
            if row_idx % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9"/>')
                cell._tc.get_or_add_tcPr().append(shading)

    doc.add_paragraph()  # Space after table


def parse_markdown_table(lines, start_idx):
    """Parse a markdown table starting at start_idx. Returns (headers, rows, end_idx)."""
    headers = []
    rows = []
    i = start_idx

    # First line should be header
    if i < len(lines) and '|' in lines[i]:
        parts = [p.strip() for p in lines[i].split('|')]
        headers = [p for p in parts if p and not p.startswith('-')]
        i += 1

    # Skip separator line (|---|---|)
    if i < len(lines) and re.match(r'^[\s|:-]+$', lines[i]):
        i += 1

    # Parse data rows
    while i < len(lines) and '|' in lines[i]:
        parts = [p.strip() for p in lines[i].split('|')]
        row = [p for p in parts if p]
        if row and not re.match(r'^[-:]+$', row[0]):
            rows.append(row)
        i += 1

    return headers, rows, i


def add_image_to_doc(doc, image_path, width_inches=6.0, caption=None):
    """Add an image to the document with optional caption."""
    try:
        doc.add_picture(str(image_path), width=Inches(width_inches))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if caption:
            cap_p = doc.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap_p.add_run(caption)
            run.font.size = Pt(10)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        return True
    except Exception as e:
        print(f"  Error adding image {image_path}: {e}")
        return False


def parse_markdown_to_doc(doc, md_content):
    """Parse markdown and add to document."""
    lines = md_content.split('\n')
    code_block = False
    code_lines = []

    i = 0
    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.startswith('```'):
            if code_block:
                # End code block
                if code_lines:
                    p = doc.add_paragraph()
                    p.style = 'No Spacing'
                    for code_line in code_lines:
                        run = p.add_run(code_line + '\n')
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
                    doc.add_paragraph()
                code_lines = []
                code_block = False
            else:
                code_block = True
            i += 1
            continue

        if code_block:
            code_lines.append(line)
            i += 1
            continue

        # Skip image references (we add images separately)
        if line.startswith('!['):
            i += 1
            continue

        # Markdown tables - detect and parse
        if '|' in line and not code_block:
            # Check if this is a table (has | and next line is separator)
            if i + 1 < len(lines) and re.match(r'^[\s|:-]+$', lines[i + 1]):
                headers, rows, new_i = parse_markdown_table(lines, i)
                if headers and rows:
                    add_styled_table(doc, headers, rows)
                i = new_i
                continue

        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            p = doc.add_paragraph()
            run = p.add_run(line[5:])
            run.font.bold = True
            run.font.size = Pt(11)
        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\. ', line):
            doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style='List Number')
        # Regular paragraph
        elif line.strip():
            # Clean markdown
            clean = line
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', clean)
            clean = re.sub(r'\*(.+?)\*', r'\1', clean)
            clean = re.sub(r'`(.+?)`', r'\1', clean)
            clean = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', clean)
            doc.add_paragraph(clean)

        i += 1


def add_markdown_file(doc, md_path):
    """Add content from a markdown file."""
    if not md_path.exists():
        print(f"  Archivo no encontrado: {md_path}")
        return

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    parse_markdown_to_doc(doc, content)


def add_diagram_section(doc, svg_path, title, description):
    """Add a diagram with title and description."""
    doc.add_heading(title, level=3)

    if description:
        p = doc.add_paragraph()
        run = p.add_run(description)
        run.font.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # Convert SVG to PNG
    png_path = SCREENSHOTS_DIR / f"{svg_path.stem}.png"
    if convert_svg_to_png(svg_path, png_path):
        add_image_to_doc(doc, png_path, width_inches=6.5)
    else:
        doc.add_paragraph(f"[Ver diagrama: {svg_path.name}]")

    doc.add_paragraph()


def add_demo_section(doc, png_path, title, description):
    """Add a demo screenshot with description."""
    doc.add_heading(title, level=2)

    if description:
        p = doc.add_paragraph()
        run = p.add_run(description)
        run.font.size = Pt(11)
        doc.add_paragraph()

    if png_path.exists():
        add_image_to_doc(doc, png_path, width_inches=6.5, caption=f"Figura: {title}")
    else:
        doc.add_paragraph(f"[Captura de demo no disponible: {png_path.name}]")

    doc.add_paragraph()


def main():
    """Generate the complete project document."""
    print("=" * 60)
    print("Generador de Documento de Proyecto Xcapit FHE-ML")
    print("=" * 60)

    setup_directories()

    doc = Document()
    doc = create_document_styles(doc)

    # 1. Cover page
    print("\n1. Generando portada...")
    add_cover_page(doc)

    # 2. Table of contents
    print("2. Generando tabla de contenidos...")
    add_table_of_contents(doc)

    # 3. Introduction
    print("3. Agregando introducción...")
    doc.add_heading("Introducción", level=1)
    intro = """
Xcapit FHE-ML es una plataforma revolucionaria que permite realizar Machine Learning sobre datos encriptados utilizando Encriptación Homomórfica Completa (FHE).

Esta tecnología resuelve uno de los mayores desafíos de la privacidad en la era del Big Data y la Inteligencia Artificial: ¿cómo podemos beneficiarnos del poder del Machine Learning sin exponer nuestros datos sensibles?

Con Xcapit FHE-ML:
"""
    doc.add_paragraph(intro.strip())
    doc.add_paragraph("Los datos del usuario NUNCA se descifran durante el procesamiento", style='List Bullet')
    doc.add_paragraph("El servidor de ML opera sobre datos encriptados sin poder ver su contenido", style='List Bullet')
    doc.add_paragraph("Solo el propietario de los datos puede descifrar los resultados", style='List Bullet')
    doc.add_paragraph("Cumplimiento automático con regulaciones como HIPAA, GDPR y LGPD", style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph("Esta documentación proporciona una guía completa desde los fundamentos teóricos de la encriptación homomórfica hasta ejemplos prácticos de implementación.")
    doc.add_page_break()

    # 4. Theory chapters
    print("4. Agregando capítulos teóricos...")
    theory_files = [
        (DOCS_DIR / "theory" / "01-homomorphic-encryption.md", "Encriptación Homomórfica"),
        (DOCS_DIR / "theory" / "02-ckks-scheme.md", "El Esquema CKKS"),
        (DOCS_DIR / "theory" / "03-ml-on-encrypted-data.md", "ML sobre Datos Encriptados"),
        (DOCS_DIR / "theory" / "04-polynomial-approximations.md", "Aproximaciones Polinomiales"),
    ]

    for md_file, title in theory_files:
        if md_file.exists():
            print(f"   - {title}")
            add_markdown_file(doc, md_file)
            doc.add_page_break()

    # 5. Architecture
    print("5. Agregando arquitectura...")
    arch_file = DOCS_DIR / "guides" / "01-architecture.md"
    if arch_file.exists():
        add_markdown_file(doc, arch_file)
        doc.add_page_break()

    # 6. ML Models
    print("6. Agregando modelos de ML...")
    models_file = DOCS_DIR / "guides" / "03-ml-models.md"
    if models_file.exists():
        add_markdown_file(doc, models_file)
        doc.add_page_break()

    # 7. Quickstart
    print("7. Agregando guía de inicio rápido...")
    quickstart_file = DOCS_DIR / "guides" / "05-quickstart.md"
    if quickstart_file.exists():
        add_markdown_file(doc, quickstart_file)
        doc.add_page_break()

    # 8. Demo Screenshots (captured with Selenium)
    print("8. Agregando demos capturados con Selenium...")
    doc.add_heading("Demos Visuales de la Plataforma", level=1)

    p = doc.add_paragraph()
    run = p.add_run(
        "Esta sección muestra capturas de pantalla reales de la plataforma Xcapit FHE-ML en funcionamiento, "
        "demostrando el flujo completo de encriptación, predicción y descifrado de datos."
    )
    run.font.size = Pt(11)
    doc.add_paragraph()

    demos = [
        ("demo_real_fhe.png", "Demostración Real de FHE en Acción",
         "Ejecución REAL del sistema Xcapit FHE-ML mostrando operaciones de Encriptación Homomórfica Completa (FHE) "
         "funcionando. Incluye: (1) Demo de encriptación CKKS con tiempos reales de creación de contexto, cifrado y "
         "descifrado, mostrando error de precisión ~10^-9. (2) Predicción de regresión lineal sobre datos encriptados "
         "con comparación entre resultado FHE y plaintext. (3) Diagnóstico médico con privacidad total usando "
         "regresión logística sobre datos de pacientes encriptados, demostrando 98.5% de accuracy en clasificación."),
        ("demo_01_dashboard.png", "Dashboard Principal de la Plataforma",
         "Vista principal del sistema Xcapit FHE-ML mostrando el monitoreo en tiempo real de predicciones "
         "encriptadas, estadísticas de uso, modelos disponibles, estado de seguridad y actividad reciente. "
         "Incluye métricas de latencia, nivel de seguridad 128-bit y badges de compliance (HIPAA, GDPR, PCI-DSS, SOC 2)."),
        ("demo_02_prediction.png", "Consola de Predicción ML en Vivo",
         "Interfaz de predicción en tiempo real mostrando el pipeline completo de inferencia FHE: desde los datos "
         "de entrada del cliente, pasando por la encriptación CKKS, hasta la inferencia en el servidor y el resultado "
         "descifrado. Incluye métricas de performance, logs de operación y garantías de privacidad."),
        ("demo_03_healthcare.png", "Plataforma Healthcare - Diagnóstico Cardiovascular",
         "Caso de uso empresarial real: sistema de diagnóstico médico HIPAA-compliant que procesa datos de pacientes (PHI) "
         "con privacidad total. Muestra el flujo de datos encriptados desde el hospital al servidor ML, con audit trail "
         "en blockchain, parámetros criptográficos y recomendaciones médicas basadas en ML."),
    ]

    for png_name, title, description in demos:
        png_path = SCREENSHOTS_DIR / png_name
        print(f"   - {title}")
        add_demo_section(doc, png_path, title, description)

    doc.add_page_break()

    # 9. SVG Diagrams
    print("9. Agregando diagramas SVG...")
    doc.add_heading("Diagramas del Sistema", level=1)

    p = doc.add_paragraph()
    run = p.add_run(
        "Diagramas técnicos que ilustran los conceptos fundamentales y la arquitectura del sistema."
    )
    run.font.size = Pt(11)
    doc.add_paragraph()

    diagrams = [
        ("fhe-overview.svg", "Visión General de FHE",
         "Flujo de datos entre cliente y servidor mostrando cómo los datos permanecen encriptados durante todo el proceso."),
        ("ckks-encoding.svg", "Codificación CKKS",
         "Proceso paso a paso de cómo los números reales se convierten en polinomios encriptados usando el esquema CKKS."),
        ("system-architecture.svg", "Arquitectura del Sistema",
         "Componentes principales de la plataforma: SDK Cliente, API Gateway, Servidor ML y Blockchain."),
        ("linear-regression-fhe.svg", "Regresión Lineal FHE",
         "Flujo detallado de una predicción de regresión lineal sobre datos encriptados con la matemática involucrada."),
        ("sigmoid-approximation.svg", "Aproximación del Sigmoid",
         "Comparación entre la función sigmoid real y su aproximación polinomial usada en FHE."),
        ("decision-tree-fhe.svg", "Árbol de Decisión FHE",
         "Cómo las comparaciones suaves (soft comparisons) permiten evaluar árboles de decisión sobre datos encriptados."),
        ("ml-pipeline-fhe.svg", "Pipeline de ML",
         "Separación entre la fase de entrenamiento (texto plano) y la fase de inferencia (encriptada)."),
        ("performance-metrics.svg", "Métricas de Performance",
         "Tiempos de ejecución por operación FHE: encriptación, regresión lineal, regresión logística y árboles de decisión."),
        ("model-accuracy.svg", "Comparación de Accuracy",
         "Comparación de precisión entre modelos en texto plano y modelos FHE, demostrando pérdida mínima de accuracy."),
        ("security-comparison.svg", "Comparación de Seguridad",
         "Análisis comparativo de FHE vs métodos tradicionales de protección de datos (TLS, encriptación en reposo)."),
    ]

    for svg_name, title, description in diagrams:
        svg_path = DIAGRAMS_DIR / svg_name
        if svg_path.exists():
            print(f"   - {title}")
            add_diagram_section(doc, svg_path, title, description)

    doc.add_page_break()

    # 10. Glossary
    print("10. Agregando glosario...")
    glossary_file = DOCS_DIR / "glossary.md"
    if glossary_file.exists():
        add_markdown_file(doc, glossary_file)
        doc.add_page_break()

    # 11. References
    print("11. Agregando referencias...")
    doc.add_heading("Referencias", level=1)

    references = [
        "Cheon, J. H., Kim, A., Kim, M., & Song, Y. (2017). Homomorphic encryption for arithmetic of approximate numbers. ASIACRYPT 2017.",
        "Gentry, C. (2009). Fully homomorphic encryption using ideal lattices. STOC 2009.",
        "Brakerski, Z., Gentry, C., & Vaikuntanathan, V. (2014). (Leveled) fully homomorphic encryption without bootstrapping. ACM TOCT.",
        "Microsoft SEAL Library. https://github.com/microsoft/SEAL",
        "TenSEAL Library. https://github.com/OpenMined/TenSEAL",
        "Gilad-Bachrach, R., et al. (2016). CryptoNets: Applying Neural Networks to Encrypted Data. ICML 2016.",
    ]

    for i, ref in enumerate(references, 1):
        doc.add_paragraph(f"[{i}] {ref}")

    # Save document
    output_path = OUTPUT_DIR / "Xcapit_FHE-ML_Documentacion_Completa.docx"
    doc.save(str(output_path))

    print("\n" + "=" * 60)
    print(f"✅ Documento generado exitosamente:")
    print(f"   {output_path}")

    # Show file size
    size_kb = output_path.stat().st_size / 1024
    print(f"   Tamaño: {size_kb:.1f} KB")
    print("=" * 60)

    return output_path


if __name__ == "__main__":
    main()
